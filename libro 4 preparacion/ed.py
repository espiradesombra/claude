from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

# Cargar el documento original
doc = Document("Cribas_cotas_y_estructuras_de_primos_VMA (1).docx")

# Crear nuevo documento con diseño editorial
design_doc = Document()

# Portada
title = design_doc.add_heading("Cribas, cotas y estructuras modulares de los primos", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = design_doc.add_paragraph("Autor: Víctor Manzanares Alberola")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
date = design_doc.add_paragraph(f"Fecha: {datetime.today().strftime('%d/%m/%Y')}")
date.alignment = WD_ALIGN_PARAGRAPH.CENTER
design_doc.add_page_break()


    

# Añadir tabla de contenido
design_doc.add_page_break()
toc_heading = design_doc.add_heading("Índice", level=1)

# Insertar campo de índice automático
paragraph = design_doc.add_paragraph()
run = paragraph.add_run()
fldChar = OxmlElement('w:fldChar')
fldChar.set(qn('w:fldCharType'), 'begin')
run._r.append(fldChar)



fldChar = OxmlElement('w:fldChar')
fldChar.set(qn('w:fldCharType'), 'separate')
run._r.append(fldChar)

fldChar = OxmlElement('w:fldChar')
fldChar.set(qn('w:fldCharType'), 'end')
run._r.append(fldChar)

# Guardar el documento con diseño editorial
design_doc.save("Cribas_cotas_y_estructuras_EDITADO.docx")
print("Documento editado con diseño editorial y portada guardado como 'Cribas_cotas_y_estructuras_EDITADO.docx'")
